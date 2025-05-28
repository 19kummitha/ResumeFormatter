import json
import re
import PyPDF2
from docx import Document
import subprocess
import platform
import tempfile
import os


def extract_text_from_pdf(file_path: str) -> str:
    """Legacy function to extract text from PDF - kept for backward compatibility"""
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files using python-docx with enhanced table parsing"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables with better structure preservation
        for table_idx, table in enumerate(doc.tables):
            text += f"\n--- EXPERIENCE TABLE {table_idx + 1} START ---\n"
            
            # Get headers if they exist
            if table.rows:
                header_row = table.rows[0]
                headers = []
                for cell in header_row.cells:
                    headers.append(cell.text.strip())
                text += "HEADERS: " + " | ".join(headers) + "\n"
                
                # Process ALL data rows - explicitly count and extract each one
                total_rows = len(table.rows) - 1  # Exclude header
                text += f"TOTAL_EXPERIENCE_ROWS: {total_rows}\n"
                
                for row_idx, row in enumerate(table.rows[1:], 1):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                        if cell_text:  # Only add non-empty cells
                            row_data.append(cell_text)
                    
                    if row_data:  # Only add rows with actual data
                        text += f"EXPERIENCE_ROW_{row_idx}: " + " | ".join(row_data) + "\n"
                        print(f"Extracted experience row {row_idx}: {' | '.join(row_data[:3])}")  # Log first 3 fields
            
            text += f"--- EXPERIENCE TABLE {table_idx + 1} END ---\n\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                text += "\n--- HEADER SECTION ---\n"
                for para in section.header.paragraphs:
                    text += para.text + "\n"
            
            # Footer  
            if section.footer:
                text += "\n--- FOOTER SECTION ---\n"
                for para in section.footer.paragraphs:
                    text += para.text + "\n"
        
        print(f"Enhanced DOCX extraction completed: {len(text)} characters with structured table data")
        return text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}")


def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert DOCX to PDF using Aspose.Words library with enhanced error handling"""
    try:
        import aspose.words as aw
        
        # Create a temporary PDF file path
        temp_dir = tempfile.gettempdir()
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(temp_dir, f"{base_name}_converted.pdf")
        
        print(f"Converting DOCX to PDF using Aspose.Words...")
        
        try:
            # Load the DOCX document
            doc = aw.Document(docx_path)
            
            # Save as PDF with high quality settings
            pdf_save_options = aw.saving.PdfSaveOptions()
            pdf_save_options.compliance = aw.saving.PdfCompliance.PDF17
            pdf_save_options.image_compression = aw.saving.PdfImageCompression.JPEG
            pdf_save_options.jpeg_quality = 90
            pdf_save_options.text_compression = aw.saving.PdfTextCompression.FLATE
            pdf_save_options.optimize_output = True
            
            # Save the document as PDF
            doc.save(pdf_path, pdf_save_options)
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"Successfully converted DOCX to PDF using Aspose.Words: {pdf_path}")
                return pdf_path
            else:
                raise Exception("PDF file was not created or is empty")
                
        except Exception as aspose_error:
            print(f"Aspose.Words conversion failed: {str(aspose_error)}")
            raise Exception(f"Aspose.Words conversion error: {str(aspose_error)}")
            
    except ImportError:
        print("Aspose.Words not available, trying alternative methods...")
        
        # Fallback to system-specific methods
        system = platform.system().lower()
        temp_dir = tempfile.gettempdir()
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(temp_dir, f"{base_name}_converted.pdf")
        
        if system == "windows":
            # Try Windows COM as fallback
            try:
                import pythoncom
                import win32com.client
                from win32com.client import constants
                
                pythoncom.CoInitialize()
                
                try:
                    word_app = win32com.client.Dispatch("Word.Application")
                    word_app.Visible = False
                    word_app.DisplayAlerts = 0
                    
                    try:
                        doc = word_app.Documents.Open(docx_path)
                        doc.ExportAsFixedFormat(
                            OutputFileName=pdf_path,
                            ExportFormat=constants.wdExportFormatPDF,
                            OpenAfterExport=False,
                            OptimizeFor=constants.wdExportOptimizeForMinSize,
                            Range=constants.wdExportAllDocument,
                            Item=constants.wdExportDocumentContents,
                            IncludeDocProps=True,
                            KeepIRM=True,
                            CreateBookmarks=constants.wdExportCreateNoBookmarks,
                            DocStructureTags=True,
                            BitmapMissingFonts=True,
                            UseDocumentImageResolution=False
                        )
                        
                        print(f"Successfully converted DOCX to PDF using Word COM: {pdf_path}")
                        return pdf_path
                        
                    finally:
                        try:
                            if 'doc' in locals():
                                doc.Close(SaveChanges=False)
                        except:
                            pass
                        
                        try:
                            word_app.Quit()
                        except:
                            pass
                            
                finally:
                    pythoncom.CoUninitialize()
                    
            except ImportError:
                print("pywin32 not available")
                raise Exception("Neither Aspose.Words nor pywin32 are available on Windows")
            except Exception as e:
                print(f"Windows COM conversion failed: {str(e)}")
                raise Exception(f"All Windows conversion methods failed: {str(e)}")
                
        elif system in ["darwin", "linux"]:  # macOS or Linux
            try:
                # Check if LibreOffice is available
                result = subprocess.run(['libreoffice', '--version'], 
                                      capture_output=True, text=True, timeout=10)
                if result.returncode != 0:
                    raise Exception("LibreOffice not found")
                
                # Convert using LibreOffice headless mode
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    docx_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0 and os.path.exists(pdf_path):
                    print(f"Successfully converted DOCX to PDF using LibreOffice: {pdf_path}")
                    return pdf_path
                else:
                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                    
            except subprocess.TimeoutExpired:
                raise Exception("LibreOffice conversion timed out")
            except FileNotFoundError:
                raise Exception("LibreOffice not found in system PATH")
            except Exception as e:
                print(f"LibreOffice conversion failed: {str(e)}")
                raise Exception(f"LibreOffice conversion failed: {str(e)}")
        else:
            raise Exception(f"Unsupported operating system: {system}")
            
    except Exception as e:
        print(f"DOCX to PDF conversion failed: {str(e)}")
        raise RuntimeError(f"Failed to convert DOCX to PDF: {str(e)}")
def clean_json_string(raw: str):
    # Remove triple backticks and language hint (```json)
    cleaned = re.sub(r"^```json\s*|\s*```$", "", raw.strip(), flags=re.MULTILINE)
    return json.loads(cleaned)